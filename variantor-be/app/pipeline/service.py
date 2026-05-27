import json
import re
import threading
import uuid
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from pypdf import PdfReader
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas
from werkzeug.datastructures import FileStorage

from app.config import Config
from app.db import connect, json_dumps
from app.errors import NotFoundError, PipelineStateError, ValidationError
from app.llm import create_provider
from app.llm.base import LLMResponse
from app.pipeline.prompts import PromptSet


TOTAL_PIPELINE_STEPS = 4
GENERATION_STEP = 3
EVALUATION_STEP = 4
SCORE_RE = re.compile(r"\b(10|[1-9])\b")
TAG_RE = re.compile(r"(?is)<[^>]+>")
SAFE_MODEL_RE = re.compile(r"^[A-Za-z0-9._:-]+$")
SUPPORTED_PROVIDERS = {"mock", "gigachat", "openai"}
PDF_FONT_NAME = "DejaVuSans"
PDF_FONT_PATHS = (
    "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    "/usr/local/share/fonts/DejaVuSans.ttf",
)
_pdf_font_registered = False


@dataclass
class TaskSection:
    number: int
    heading: str
    html: str
    start: int
    end: int


@dataclass
class SavedFile:
    relative_path: str
    mime_type: str
    size_bytes: int


@dataclass(frozen=True)
class LLMSelection:
    provider: str
    model: str


class PipelineService:
    def __init__(self, config: Config):
        self.config = config
        self.prompts = PromptSet(config.PROMPTS_DIR)
        self._providers: dict[str, Any] = {}
        Path(config.FILE_STORAGE_DIR).mkdir(parents=True, exist_ok=True)

    def llm_options(self) -> dict:
        return {
            "default_provider": self.config.LLM_PROVIDER,
            "providers": provider_options(self.config),
        }

    def _provider(self, selection: LLMSelection):
        if selection.provider not in self._providers:
            self._providers[selection.provider] = create_provider(self.config, selection.provider)
        return self._providers[selection.provider]

    def create_assignment(self, title: str) -> dict:
        with connect() as conn:
            row = conn.execute(
                """
                insert into assignments (title, subject, variant_count, status, user_id)
                values (nullif(%s, ''), 'english', 1, 'created', 'default-user')
                returning *
                """,
                (title.strip(),),
            ).fetchone()
        return self.assignment_payload(row)

    def get_assignment(self, assignment_id: str) -> dict:
        assignment = self._assignment_or_404(assignment_id)
        images = self._images_by_assignment_id(assignment_id, missing_ok=True)
        latest_run = self._latest_run_for_assignment(assignment_id)
        return self.assignment_payload(assignment, image=(images[-1] if images else None), images=images, latest_run=latest_run)

    def save_assignment_image(self, assignment_id: str, file: FileStorage | None) -> dict:
        self._assignment_or_404(assignment_id)
        if file is None or not file.filename:
            raise ValidationError("Image file is required.", code="validation_error")
        saved = self._save_file(assignment_id, file)
        with connect() as conn:
            image = conn.execute(
                """
                insert into assignment_images (
                    assignment_id, original_filename, stored_path, mime_type, size_bytes
                )
                values (%s, nullif(%s, ''), %s, %s, %s)
                returning *
                """,
                (assignment_id, file.filename, saved.relative_path, saved.mime_type, saved.size_bytes),
            ).fetchone()
            conn.execute("update assignments set status = 'image_uploaded', updated_at = now() where id = %s", (assignment_id,))
        return {
            "assignment_id": assignment_id,
            "image": self.image_payload(image),
            "status": "image_uploaded",
        }

    def save_assignment_files(self, assignment_id: str, files: list[FileStorage]) -> dict:
        self._assignment_or_404(assignment_id)
        valid_files = [item for item in files if item and item.filename]
        if not valid_files:
            raise ValidationError("At least one file is required.", code="validation_error")
        saved_files = [self._save_file(assignment_id, file) for file in valid_files]
        rows = []
        with connect() as conn:
            conn.execute("delete from assignment_images where assignment_id = %s", (assignment_id,))
            for index, (file, saved) in enumerate(zip(valid_files, saved_files)):
                row = conn.execute(
                    """
                    insert into assignment_images (
                        assignment_id, original_filename, stored_path, mime_type, size_bytes, position
                    )
                    values (%s, nullif(%s, ''), %s, %s, %s, %s)
                    returning *
                    """,
                    (assignment_id, file.filename, saved.relative_path, saved.mime_type, saved.size_bytes, index + 1),
                ).fetchone()
                rows.append(row)
            conn.execute("update assignments set status = 'image_uploaded', updated_at = now() where id = %s", (assignment_id,))
        return {
            "assignment_id": assignment_id,
            "files": [self.image_payload(row) for row in rows],
            "status": "image_uploaded",
        }

    def preview_pdf_for_file(self, file: FileStorage | None) -> bytes:
        if file is None or not file.filename:
            raise ValidationError("File is required.", code="validation_error")
        ext = Path(file.filename).suffix.lower()
        raw = file.read() or b""
        file.stream.seek(0)
        if not raw:
            raise ValidationError("File is empty.", code="validation_error")
        if ext == ".pdf":
            return raw
        if ext == ".docx":
            text = read_docx_text_from_bytes(raw)
            return text_to_pdf_bytes(text)
        if ext == ".doc":
            text = raw.decode("utf-8", errors="ignore")
            return text_to_pdf_bytes(text)
        raise ValidationError("Preview is available only for PDF, DOC and DOCX files.", code="preview_not_supported")

    def start_extraction(self, assignment_id: str, options: dict[str, Any] | None = None) -> dict:
        options = options or {}
        self._assignment_or_404(assignment_id)
        use_default_source = bool(options.get("use_default_source"))
        if not use_default_source:
            self._images_by_assignment_id(assignment_id)

        step_selection = resolve_llm_selection(
            options,
            provider_key="step_provider",
            model_key="step_model",
            config=self.config,
        )
        with connect() as conn:
            run = conn.execute(
                """
                insert into extraction_runs (
                    assignment_id, status, current_step, provider, model, prompt_version,
                    step_results, parsed_content, warnings
                )
                values (%s, 'pending', 1, %s, %s, %s, '[]'::jsonb, null, '[]'::jsonb)
                returning *
                """,
                (
                    assignment_id,
                    step_selection.provider,
                    step_selection.model,
                    self.prompts.version,
                ),
            ).fetchone()
            conn.execute("update assignments set status = 'extracting', updated_at = now() where id = %s", (assignment_id,))

        self._background(self.execute_step, str(run["id"]), 1, step_selection, 1, use_default_source, step_selection)
        return {"extraction_run_id": str(run["id"]), "status": run["status"]}

    def get_run(self, run_id: str) -> dict:
        return self.run_payload(self._run_or_404(run_id))

    def continue_run(self, run_id: str, options: dict[str, Any] | None = None) -> dict:
        options = options or {}
        run = self._run_or_404(run_id)
        if run["status"] != "awaiting_confirmation" or run["current_step"] >= TOTAL_PIPELINE_STEPS:
            raise PipelineStateError("Pipeline is not waiting for confirmation.", code="pipeline_not_ready")

        next_step = int(run["current_step"]) + 1
        step_selection = resolve_llm_selection(
            options,
            provider_key="step_provider",
            model_key="step_model",
            config=self.config,
        )
        if next_step == GENERATION_STEP and ("final_provider" in options or "final_model" in options):
            step_selection = resolve_llm_selection(
                options,
                provider_key="final_provider",
                model_key="final_model",
                config=self.config,
                fallback=step_selection,
            )
        evaluation_selection = resolve_llm_selection(
            options,
            provider_key="evaluation_provider",
            model_key="evaluation_model",
            config=self.config,
            fallback=step_selection,
        )
        variant_count = resolve_variant_count(options.get("variant_count"))
        self._mark_step_running(run_id, next_step)
        self._background(self.execute_step, run_id, next_step, step_selection, variant_count, False, evaluation_selection)
        return {"extraction_run_id": run_id, "status": "running", "next_step": next_step}

    def update_step(self, run_id: str, step: int, content: str) -> dict:
        if step < 1 or step > TOTAL_PIPELINE_STEPS or not content.strip():
            raise ValidationError("Step or content is invalid.", code="invalid_step")
        run = self._run_or_404(run_id)
        if run["status"] in {"pending", "running"}:
            raise PipelineStateError("Pipeline is not ready for this action.", code="pipeline_not_ready")

        previous_results = parse_results(run["step_results"])
        results = keep_results_before(previous_results, step)
        results.append(step_result(step, pipeline_step_key(step), pipeline_step_title(step), content.strip()))
        if step == GENERATION_STEP and run["status"] == "succeeded":
            score = result_content(previous_results, "self_score")
            if score:
                results.append(step_result(EVALUATION_STEP, "self_score", pipeline_step_title(EVALUATION_STEP), score))

        status = "awaiting_confirmation"
        assignment_status = "processing_waiting"
        current_step = step
        if step == TOTAL_PIPELINE_STEPS or (step == GENERATION_STEP and run["status"] == "succeeded"):
            status = "succeeded"
            assignment_status = "processed"
            current_step = TOTAL_PIPELINE_STEPS

        parsed = build_pipeline_content(results)
        existing = run["parsed_content"] or {}
        if step == GENERATION_STEP and run["status"] == "succeeded" and existing.get("variants_html"):
            parsed["variants_html"] = existing["variants_html"]
            parsed["variant_html"] = content.strip()
            parsed["selected_variant"] = 1
            for index, variant in enumerate(existing["variants_html"]):
                if variant.strip() == content.strip():
                    parsed["selected_variant"] = index + 1
                    break
            if existing.get("answers_by_variant"):
                parsed["answers_by_variant"] = existing["answers_by_variant"]
            if existing.get("answers_all"):
                parsed["answers_all"] = existing["answers_all"]

        updated = self._update_run_results(
            run_id,
            status=status,
            assignment_status=assignment_status,
            current_step=current_step,
            step_results=results,
            parsed_content=parsed,
        )
        return self.run_payload(updated)

    def regenerate_step(self, run_id: str, step: int, options: dict[str, Any] | None = None) -> dict:
        options = options or {}
        if step < 1 or step > TOTAL_PIPELINE_STEPS:
            raise ValidationError("Step is invalid.", code="invalid_step")
        run = self._run_or_404(run_id)
        if run["status"] in {"pending", "running"}:
            raise PipelineStateError("Pipeline is not ready for this action.", code="pipeline_not_ready")

        results = keep_results_before(parse_results(run["step_results"]), step)
        parsed = build_pipeline_content(results)
        self._update_run_results(
            run_id,
            status="running",
            assignment_status="extracting",
            current_step=step,
            step_results=results,
            parsed_content=parsed,
        )
        step_selection = resolve_llm_selection(
            options,
            provider_key="step_provider",
            model_key="step_model",
            config=self.config,
        )
        evaluation_selection = resolve_llm_selection(
            options,
            provider_key="evaluation_provider",
            model_key="evaluation_model",
            config=self.config,
            fallback=step_selection,
        )
        self._background(self.execute_step, run_id, step, step_selection, 1, False, evaluation_selection)
        return {"extraction_run_id": run_id, "status": "running", "step": step}

    def execute_step(
        self,
        run_id: str,
        step: int,
        selection: LLMSelection,
        variant_count: int,
        use_default_source: bool,
        evaluation_selection: LLMSelection,
    ) -> None:
        try:
            run = self._run_or_404(run_id)
            self._mark_step_running(run_id, step)
            results = parse_results(run["step_results"])
            response, result, variants, llm_input = self._run_step(
                run,
                step,
                results,
                selection,
                variant_count,
                use_default_source,
            )

            evaluation_result = None
            if step == GENERATION_STEP:
                response, result, variants, llm_input, evaluation_result = self._evaluate_and_maybe_retry(
                    run,
                    results,
                    response,
                    result,
                    variants,
                    llm_input,
                    selection,
                    variant_count,
                    evaluation_selection,
                )

            results.append(result)
            parsed = build_pipeline_content(results)
            if use_default_source and step == 1:
                parsed["used_default_html"] = True
            if (run["parsed_content"] or {}).get("used_default_html"):
                parsed["used_default_html"] = True
            if step == GENERATION_STEP and variants:
                parsed["variants_html"] = variants
                parsed["selected_variant"] = 1
                parsed["variant_html"] = variants[0]
            if evaluation_result is not None:
                results.append(evaluation_result)
                parsed["self_score"] = evaluation_result["content"]
                parsed["steps"] = results
            if step == GENERATION_STEP and variants:
                answers_by_variant, answers_all = self._generate_answers_for_variants(run, variants, evaluation_selection)
                parsed["answers_by_variant"] = answers_by_variant
                parsed["answers_all"] = answers_all

            self._insert_llm_run(
                task_type=pipeline_task_type(step),
                response=response,
                llm_input=llm_input,
                parsed_output=result,
                status="succeeded",
            )

            status = "awaiting_confirmation"
            assignment_status = "processing_waiting"
            if step in {TOTAL_PIPELINE_STEPS, GENERATION_STEP}:
                status = "succeeded"
                assignment_status = "processed"
            self._finish_step(
                run_id,
                status=status,
                assignment_status=assignment_status,
                provider=response.provider,
                model=response.model,
                raw_response=response.raw_response,
                current_step=finished_pipeline_step(step),
                step_results=results,
                parsed_content=parsed,
            )
        except Exception as exc:
            self._fail_run(run_id, selection.provider, selection.model, "", str(exc))

    def _run_step(
        self,
        run: dict,
        step: int,
        results: list[dict],
        selection: LLMSelection,
        variant_count: int,
        use_default_source: bool,
    ) -> tuple[LLMResponse, dict, list[str], dict]:
        llm_input: dict[str, Any] = {
            "assignment_id": str(run["assignment_id"]),
            "run_id": str(run["id"]),
            "step": step,
            "prompt_version": self.prompts.version,
            "provider": selection.provider,
            "model": selection.model,
        }
        if step == 1:
            if use_default_source:
                source_html = normalize_source_document(self.prompts.default_source, self.prompts.is_json_pipeline)
                llm_input["use_default_source"] = True
                return (
                    LLMResponse(source_html, source_html, "local", "default-html"),
                    step_result(1, "source_html", pipeline_step_title(1), source_html),
                    [],
                    llm_input,
                )
            provider = self._provider(selection)
            prompt = self.prompts.source_from_image_prompt()
            files = self._images_by_assignment_id(str(run["assignment_id"]))
            llm_input.update(
                {
                    "file_count": len(files),
                    "files": [{"path": item["stored_path"], "mime_type": item["mime_type"]} for item in files],
                    "prompt": prompt,
                }
            )
            parts = []
            raw_outputs = []
            last_response: LLMResponse | None = None
            for index, file in enumerate(files):
                source = self._source_html_for_file(
                    provider=provider,
                    prompt=prompt,
                    file_path=str(Path(self.config.FILE_STORAGE_DIR) / file["stored_path"]),
                    mime_type=file["mime_type"],
                    model=selection.model or None,
                    index=index + 1,
                    json_pipeline=self.prompts.is_json_pipeline,
                )
                parts.append(source["html"])
                raw_outputs.append(source["raw"])
                last_response = source["response"]
            content = merge_json_documents(parts) if self.prompts.is_json_pipeline else merge_source_documents(parts)
            response = LLMResponse(
                content=content,
                raw_response=json.dumps(raw_outputs, ensure_ascii=False, indent=2),
                provider=(last_response.provider if last_response else selection.provider),
                model=(last_response.model if last_response else selection.model),
            )
            return response, step_result(1, "source_html", pipeline_step_title(1), content), [], llm_input

        if step == 2:
            provider = self._provider(selection)
            source_html = result_content(results, "source_html")
            if not source_html:
                raise PipelineStateError("step 1 result is missing")
            sections = extract_task_sections(source_html)
            tasks = []
            raw_outputs = []
            last_response: LLMResponse | None = None
            for section in sections:
                prompt = self.prompts.parameters_prompt(section.html)
                response = provider.complete_text(prompt, model=selection.model or None)
                params = parse_task_parameters(normalize_llm_text(response.content))
                params["task_number"] = section.number
                params["heading"] = section.heading
                tasks.append(params)
                raw_outputs.append(
                    {
                        "task_number": section.number,
                        "heading": section.heading,
                        "content": normalize_llm_text(response.content),
                        "provider": response.provider,
                        "model": response.model,
                    }
                )
                last_response = response
            serialized = json.dumps({"tasks": tasks}, ensure_ascii=False, indent=2)
            llm_input.update({"task_count": len(sections), "tasks": [section.__dict__ for section in sections]})
            response = LLMResponse(
                content=serialized,
                raw_response=json.dumps(raw_outputs, ensure_ascii=False, indent=2),
                provider=(last_response.provider if last_response else "unknown"),
                model=(last_response.model if last_response else selection.model),
            )
            return response, step_result(2, "parameters", pipeline_step_title(2), serialized), [], llm_input

        if step == 3:
            provider = self._provider(selection)
            source_html = result_content(results, "source_html")
            parameters = result_content(results, "parameters")
            if not source_html or not parameters:
                raise PipelineStateError("previous pipeline results are missing")
            sections = extract_task_sections(source_html)
            bundle = parse_parameter_bundle(parameters)
            tasks = bundle["tasks"]
            if len(tasks) != len(sections):
                raise PipelineStateError(f"expected {len(sections)} task parameter sets, got {len(tasks)}")
            full_variants = []
            raw_outputs = []
            last_response: LLMResponse | None = None
            for variant_index in range(variant_count):
                generated_tasks = []
                task_outputs = []
                for index, section in enumerate(sections):
                    params = tasks[index]
                    prompt = self.prompts.generation_prompt(section.html, params, bundle.get("user_comment", ""))
                    response = provider.complete_text(prompt, model=selection.model or None)
                    task_html = normalize_llm_text(response.content)
                    if not self.prompts.is_json_pipeline and not contains_tag(task_html, "h2"):
                        task_html = ensure_section_heading(section.html, task_html)
                    generated_tasks.append(task_html)
                    task_outputs.append(
                        {
                            "task_number": params.get("task_number"),
                            "heading": params.get("heading"),
                            "content": task_html,
                            "provider": response.provider,
                            "model": response.model,
                        }
                    )
                    last_response = response
                full_variants.append(merge_task_sections(source_html, sections, generated_tasks))
                raw_outputs.append({"variant_index": variant_index + 1, "tasks": task_outputs})
            first_variant = full_variants[0]
            response = LLMResponse(
                content=first_variant,
                raw_response=json.dumps(raw_outputs, ensure_ascii=False, indent=2),
                provider=(last_response.provider if last_response else "unknown"),
                model=(last_response.model if last_response else selection.model),
            )
            llm_input.update(
                {
                    "variant_count": variant_count,
                    "task_count": len(sections),
                    "tasks": tasks,
                    "user_comment": bundle.get("user_comment", ""),
                }
            )
            return response, step_result(3, "variant_html", pipeline_step_title(3), first_variant), full_variants, llm_input

        if step == 4:
            provider = self._provider(selection)
            source_html = result_content(results, "source_html")
            variant_html = result_content(results, "variant_html")
            if not source_html or not variant_html:
                raise PipelineStateError("previous pipeline results are missing")
            prompt = self.prompts.self_evaluation_prompt(source_html, variant_html)
            llm_input["prompt"] = prompt
            response = provider.complete_text(prompt, model=selection.model or None)
            score = str(extract_score(response.content))
            return response, step_result(4, "self_score", pipeline_step_title(4), score), [], llm_input

        raise ValidationError("Unsupported pipeline step.", code="invalid_step")

    def _evaluate_and_maybe_retry(
        self,
        run: dict,
        previous_results: list[dict],
        response: LLMResponse,
        result: dict,
        variants: list[str],
        llm_input: dict,
        selection: LLMSelection,
        variant_count: int,
        evaluation_selection: LLMSelection,
    ) -> tuple[LLMResponse, dict, list[str], dict, dict | None]:
        source_html = result_content(previous_results, "source_html")
        try:
            eval_response, eval_result, score, eval_input = self._evaluate_variant(
                run,
                source_html,
                result["content"],
                evaluation_selection,
            )
            self._insert_llm_run(pipeline_task_type(EVALUATION_STEP), eval_response, eval_input, eval_result, "succeeded")
        except Exception:
            return response, result, variants, llm_input, None

        if score > 6:
            return response, result, variants, llm_input, eval_result

        try:
            retry_response, retry_result, retry_variants, retry_input = self._run_step(
                run,
                GENERATION_STEP,
                previous_results,
                selection,
                variant_count,
                False,
            )
            retry_eval_response, retry_eval_result, _, retry_eval_input = self._evaluate_variant(
                run,
                source_html,
                retry_result["content"],
                evaluation_selection,
            )
            self._insert_llm_run(pipeline_task_type(EVALUATION_STEP), retry_eval_response, retry_eval_input, retry_eval_result, "succeeded")
            return retry_response, retry_result, retry_variants, retry_input, retry_eval_result
        except Exception:
            return response, result, variants, llm_input, eval_result

    def _evaluate_variant(
        self,
        run: dict,
        source_html: str,
        variant_html: str,
        selection: LLMSelection,
    ) -> tuple[LLMResponse, dict, int, dict]:
        provider = self._provider(selection)
        prompt = self.prompts.self_evaluation_prompt(source_html, variant_html)
        llm_input = {
            "assignment_id": str(run["assignment_id"]),
            "run_id": str(run["id"]),
            "step": EVALUATION_STEP,
            "prompt_version": self.prompts.version,
            "provider": selection.provider,
            "model": selection.model,
            "prompt": prompt,
        }
        response = provider.complete_text(prompt, model=selection.model or None)
        score = extract_score(response.content)
        return response, step_result(EVALUATION_STEP, "self_score", pipeline_step_title(EVALUATION_STEP), str(score)), score, llm_input

    def _generate_answers_for_variants(
        self,
        run: dict,
        variants: list[str],
        selection: LLMSelection,
    ) -> tuple[list[str], str]:
        provider = self._provider(selection)
        answers_by_variant: list[str] = []
        for variant in variants:
            prompt = self.prompts.answer_generation_prompt(variant)
            response = provider.complete_text(prompt, model=selection.model or None)
            answers_by_variant.append(normalize_llm_text(response.content))

        all_answers = "\n".join(
            f"<section><h2>Ответы к варианту {index + 1}</h2>{answers}</section>"
            for index, answers in enumerate(answers_by_variant)
        )
        return answers_by_variant, all_answers

    def _assignment_or_404(self, assignment_id: str) -> dict:
        with connect() as conn:
            row = conn.execute("select * from assignments where id = %s", (assignment_id,)).fetchone()
        if row is None:
            raise NotFoundError("Assignment was not found.", code="assignment_not_found")
        return row

    def _run_or_404(self, run_id: str) -> dict:
        with connect() as conn:
            row = conn.execute("select * from extraction_runs where id = %s", (run_id,)).fetchone()
        if row is None:
            raise NotFoundError("Extraction run was not found.", code="extraction_not_found")
        return row

    def _image_by_assignment_id(self, assignment_id: str, *, missing_ok: bool = False) -> dict | None:
        rows = self._images_by_assignment_id(assignment_id, missing_ok=missing_ok)
        return rows[-1] if rows else None

    def _images_by_assignment_id(self, assignment_id: str, *, missing_ok: bool = False) -> list[dict]:
        with connect() as conn:
            rows = conn.execute(
                "select * from assignment_images where assignment_id = %s order by position asc, created_at asc",
                (assignment_id,),
            ).fetchall()
        if not rows and not missing_ok:
            raise NotFoundError("Assignment image was not found.", code="image_not_found")
        return rows or []

    def _latest_run_for_assignment(self, assignment_id: str) -> dict | None:
        with connect() as conn:
            return conn.execute(
                "select * from extraction_runs where assignment_id = %s order by created_at desc limit 1",
                (assignment_id,),
            ).fetchone()

    def _mark_step_running(self, run_id: str, step: int) -> None:
        with connect() as conn:
            row = conn.execute(
                """
                update extraction_runs
                set status = 'running',
                    current_step = %s,
                    started_at = coalesce(started_at, now()),
                    finished_at = null,
                    error_message = null
                where id = %s
                returning id
                """,
                (step, run_id),
            ).fetchone()
        if row is None:
            raise NotFoundError("Extraction run was not found.", code="extraction_not_found")

    def _update_run_results(
        self,
        run_id: str,
        *,
        status: str,
        assignment_status: str,
        current_step: int,
        step_results: list[dict],
        parsed_content: dict,
    ) -> dict:
        with connect() as conn:
            updated = conn.execute(
                """
                update extraction_runs
                set status = %s,
                    current_step = %s,
                    step_results = %s::jsonb,
                    parsed_content = %s::jsonb,
                    error_message = null,
                    finished_at = case when %s = 'succeeded' then now() else null end
                where id = %s
                returning *
                """,
                (status, current_step, json_dumps(step_results), json_dumps(parsed_content), status, run_id),
            ).fetchone()
            if updated is None:
                raise NotFoundError("Extraction run was not found.", code="extraction_not_found")
            conn.execute(
                "update assignments set status = %s, updated_at = now() where id = %s",
                (assignment_status, updated["assignment_id"]),
            )
        return updated

    def _finish_step(
        self,
        run_id: str,
        *,
        status: str,
        assignment_status: str,
        provider: str,
        model: str,
        raw_response: str,
        current_step: int,
        step_results: list[dict],
        parsed_content: dict,
    ) -> None:
        with connect() as conn:
            updated = conn.execute(
                """
                update extraction_runs
                set status = %s,
                    provider = nullif(%s, ''),
                    model = nullif(%s, ''),
                    prompt_version = %s,
                    raw_response = nullif(%s, ''),
                    current_step = %s,
                    step_results = %s::jsonb,
                    parsed_content = %s::jsonb,
                    warnings = '[]'::jsonb,
                    error_message = null,
                    finished_at = case when %s = 'succeeded' then now() else null end
                where id = %s
                returning *
                """,
                (
                    status,
                    provider,
                    model,
                    self.prompts.version,
                    raw_response,
                    current_step,
                    json_dumps(step_results),
                    json_dumps(parsed_content),
                    status,
                    run_id,
                ),
            ).fetchone()
            if updated is None:
                raise NotFoundError("Extraction run was not found.", code="extraction_not_found")
            conn.execute(
                "update assignments set status = %s, updated_at = now() where id = %s",
                (assignment_status, updated["assignment_id"]),
            )

    def _fail_run(self, run_id: str, provider: str, model: str, raw_response: str, message: str) -> None:
        try:
            with connect() as conn:
                updated = conn.execute(
                    """
                    update extraction_runs
                    set status = 'failed',
                        provider = nullif(%s, ''),
                        model = nullif(%s, ''),
                        prompt_version = %s,
                        raw_response = nullif(%s, ''),
                        error_message = %s,
                        finished_at = now()
                    where id = %s
                    returning *
                    """,
                    (provider, model, self.prompts.version, raw_response, message, run_id),
                ).fetchone()
                if updated is not None:
                    conn.execute(
                        "update assignments set status = 'extraction_failed', updated_at = now() where id = %s",
                        (updated["assignment_id"],),
                    )
        except Exception:
            pass

    def _insert_llm_run(
        self,
        task_type: str,
        response: LLMResponse,
        llm_input: dict,
        parsed_output: dict,
        status: str,
        error_message: str = "",
    ) -> None:
        try:
            with connect() as conn:
                conn.execute(
                    """
                    insert into llm_runs (
                        task_type, provider, model, prompt_version, input, raw_output,
                        parsed_output, status, error_message, finished_at
                    )
                    values (%s, %s, nullif(%s, ''), %s, %s::jsonb, nullif(%s, ''),
                        %s::jsonb, %s, nullif(%s, ''), now())
                    """,
                    (
                        task_type,
                        response.provider or "unknown",
                        response.model or "",
                        self.prompts.version,
                        json_dumps(llm_input),
                        response.raw_response or response.content,
                        json_dumps(parsed_output),
                        status,
                        error_message,
                    ),
                )
        except Exception:
            pass

    def _save_file(self, assignment_id: str, file: FileStorage) -> SavedFile:
        ext = Path(file.filename or "").suffix.lower()
        if ext and ext not in self.config.ALLOWED_UPLOAD_EXTENSIONS:
            raise ValidationError("File extension is not allowed.", code="invalid_file_type")
        file.stream.seek(0)
        head = file.stream.read(512)
        file.stream.seek(0)
        mime_type, ext = detect_supported_file(file.filename or "", head)
        if mime_type is None:
            raise ValidationError("Supported files: PNG, JPG, WEBP, PDF, DOC, DOCX.", code="invalid_file_type")
        relative_path = Path("assignments") / assignment_id / f"{uuid.uuid4().hex}{ext}"
        full_path = Path(self.config.FILE_STORAGE_DIR) / relative_path
        full_path.parent.mkdir(parents=True, exist_ok=True)
        temp_path = full_path.with_suffix(full_path.suffix + ".tmp")
        file.save(temp_path)
        size = temp_path.stat().st_size
        if size > self.config.MAX_UPLOAD_BYTES:
            temp_path.unlink(missing_ok=True)
            raise ValidationError("File must be 10 MB or smaller.", code="file_too_large", status_code=413)
        temp_path.replace(full_path)
        return SavedFile(relative_path=relative_path.as_posix(), mime_type=mime_type, size_bytes=size)

    def _source_html_for_file(
        self,
        provider,
        prompt: str,
        file_path: str,
        mime_type: str,
        model: str | None,
        index: int,
        json_pipeline: bool,
    ) -> dict:
        if mime_type.startswith("image/"):
            response = provider.complete_with_file(prompt, file_path=file_path, mime_type=mime_type, model=model)
            source = normalize_source_document(response.content, json_pipeline)
            return {"html": source, "raw": {"index": index, "mime_type": mime_type}, "response": response}
        if mime_type == "application/pdf":
            source = json_from_pdf(file_path, index) if json_pipeline else html_from_pdf(file_path, index)
            response = LLMResponse(content=source, raw_response=source, provider="local", model="pdf-reader")
            return {"html": source, "raw": {"index": index, "mime_type": mime_type}, "response": response}
        if mime_type in {"application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"}:
            source = json_from_doc(file_path, index) if json_pipeline else html_from_doc(file_path, index)
            response = LLMResponse(content=source, raw_response=source, provider="local", model="doc-reader")
            return {"html": source, "raw": {"index": index, "mime_type": mime_type}, "response": response}
        raise ValidationError("Unsupported file type for parsing.", code="invalid_file_type")

    def _background(self, target, *args) -> None:
        thread = threading.Thread(target=target, args=args, daemon=True)
        thread.start()

    def assignment_payload(
        self,
        assignment: dict,
        *,
        image: dict | None = None,
        images: list[dict] | None = None,
        latest_run: dict | None = None,
    ) -> dict:
        return {
            "id": str(assignment["id"]),
            "title": assignment.get("title") or "",
            "status": assignment["status"],
            "created_at": iso(assignment["created_at"]),
            "image": self.image_payload(image) if image else None,
            "files": [self.image_payload(item) for item in (images or [])],
            "latest_extraction_run": (
                {"id": str(latest_run["id"]), "status": latest_run["status"]} if latest_run else None
            ),
        }

    def image_payload(self, image: dict) -> dict:
        return {
            "id": str(image["id"]),
            "url": self.assignment_image_url(str(image["id"])),
            "preview_url": self.assignment_image_preview_url(str(image["id"])),
            "mime_type": image["mime_type"],
            "size_bytes": image["size_bytes"],
        }

    def run_payload(self, run: dict) -> dict:
        return {
            "id": str(run["id"]),
            "assignment_id": str(run["assignment_id"]),
            "status": run["status"],
            "current_step": run["current_step"],
            "provider": run.get("provider") or "",
            "model": run.get("model") or "",
            "prompt_version": run.get("prompt_version") or "",
            "parsed_content": run.get("parsed_content"),
            "step_results": run.get("step_results") or [],
            "warnings": run.get("warnings") or [],
            "error_message": run.get("error_message") or "",
            "created_at": iso(run["created_at"]),
            "started_at": iso(run.get("started_at")),
            "finished_at": iso(run.get("finished_at")),
        }

    def assignment_image_url(self, image_id: str) -> str:
        return f"{self.config.PUBLIC_FILE_BASE_URL}/assignment-images/{image_id}"

    def assignment_image_preview_url(self, image_id: str) -> str:
        return f"{self.config.PUBLIC_FILE_BASE_URL}/assignment-images/{image_id}/preview.pdf"

    def image_full_path(self, image: dict) -> Path:
        return Path(self.config.FILE_STORAGE_DIR) / image["stored_path"]

    def image_by_id_or_404(self, image_id: str) -> dict:
        with connect() as conn:
            row = conn.execute("select * from assignment_images where id = %s", (image_id,)).fetchone()
        if row is None:
            raise NotFoundError("Assignment image was not found.", code="image_not_found")
        return row

    def image_preview_pdf_bytes(self, image: dict) -> bytes:
        path = self.image_full_path(image)
        mime_type = str(image.get("mime_type") or "")
        if mime_type == "application/pdf":
            return path.read_bytes()
        if mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            text = read_docx_text(path)
            return text_to_pdf_bytes(text)
        if mime_type == "application/msword":
            text = read_doc_text(path)
            return text_to_pdf_bytes(text)
        raise ValidationError("Preview PDF is available only for PDF, DOC and DOCX files.", code="preview_not_supported")


def detect_supported_file(filename: str, head: bytes) -> tuple[str | None, str]:
    ext = Path(filename).suffix.lower()
    if head.startswith(b"\xff\xd8\xff"):
        return "image/jpeg", ".jpg"
    if head.startswith(b"\x89PNG\r\n\x1a\n"):
        return "image/png", ".png"
    if len(head) >= 12 and head[:4] == b"RIFF" and head[8:12] == b"WEBP":
        return "image/webp", ".webp"
    if head.startswith(b"%PDF"):
        return "application/pdf", ".pdf"
    if ext == ".docx":
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document", ".docx"
    if ext == ".doc":
        return "application/msword", ".doc"
    return None, ""


def html_from_pdf(file_path: str, index: int) -> str:
    pages = []
    reader = PdfReader(file_path)
    for page in reader.pages:
        text = (page.extract_text() or "").strip()
        if text:
            pages.append(text)
    body = "\n\n".join(pages).strip()
    return text_to_html_document(body, index)


def json_from_pdf(file_path: str, index: int) -> str:
    pages = []
    reader = PdfReader(file_path)
    for page in reader.pages:
        text = (page.extract_text() or "").strip()
        if text:
            pages.append(text)
    return text_to_json_document("\n\n".join(pages).strip(), index)


def html_from_doc(file_path: str, index: int) -> str:
    suffix = Path(file_path).suffix.lower()
    if suffix == ".docx":
        return text_to_html_document(read_docx_text(Path(file_path)), index)
    return text_to_html_document(read_doc_text(Path(file_path)), index)


def json_from_doc(file_path: str, index: int) -> str:
    suffix = Path(file_path).suffix.lower()
    if suffix == ".docx":
        return text_to_json_document(read_docx_text(Path(file_path)), index)
    return text_to_json_document(read_doc_text(Path(file_path)), index)


def text_to_html_document(text: str, index: int) -> str:
    clean = text.strip() or "Не удалось извлечь текст из файла."
    paragraphs = "".join(f"<p>{escape_html(line)}</p>" for line in clean.splitlines() if line.strip())
    if not paragraphs:
        paragraphs = f"<p>{escape_html(clean)}</p>"
    return (
        "<!doctype html><html><body>"
        f"<h2>Задание {index}</h2>"
        f"{paragraphs}"
        "</body></html>"
    )

def read_docx_text(path: Path) -> str:
    document = Document(str(path))
    lines = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    return "\n\n".join(lines)

def read_docx_text_from_bytes(raw: bytes) -> str:
    document = Document(BytesIO(raw))
    lines = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    return "\n\n".join(lines)


def read_doc_text(path: Path) -> str:
    with path.open("rb") as f:
        raw = f.read()
    return raw.decode("utf-8", errors="ignore")


def text_to_pdf_bytes(text: str) -> bytes:
    font_name = pdf_font_name()
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    pdf.setFont(font_name, 10)
    width, height = A4
    y = height - 40
    for raw_line in (text.strip() or "Документ пуст.").splitlines():
        line = raw_line.strip() or " "
        while len(line) > 110:
            pdf.drawString(40, y, line[:110])
            y -= 14
            line = line[110:]
            if y < 40:
                pdf.showPage()
                pdf.setFont(font_name, 10)
                y = height - 40
        pdf.drawString(40, y, line)
        y -= 14
        if y < 40:
            pdf.showPage()
            pdf.setFont(font_name, 10)
            y = height - 40
    pdf.save()
    return buffer.getvalue()


def pdf_font_name() -> str:
    global _pdf_font_registered
    if _pdf_font_registered:
        return PDF_FONT_NAME
    for font_path in PDF_FONT_PATHS:
        if Path(font_path).exists():
            pdfmetrics.registerFont(TTFont(PDF_FONT_NAME, font_path))
            _pdf_font_registered = True
            return PDF_FONT_NAME
    return "Helvetica"


def text_to_json_document(text: str, index: int) -> str:
    clean = text.strip() or "Не удалось извлечь текст из файла."
    task = {
        "title": f"Задание {index}",
        "text": "".join(f"<p>{escape_html(line)}</p>" for line in clean.splitlines() if line.strip()),
    }
    return json.dumps({"tasks": [task]}, ensure_ascii=False, indent=2)


def normalize_source_document(content: str, json_pipeline: bool) -> str:
    normalized = normalize_llm_text(content)
    if not json_pipeline:
        return normalized
    return json.dumps(parse_json_document(normalized), ensure_ascii=False, indent=2)


def parse_json_document(content: str) -> dict:
    parsed = json.loads(normalize_llm_text(content))
    if isinstance(parsed, list):
        return {"tasks": parsed}
    if isinstance(parsed, dict) and isinstance(parsed.get("tasks"), list):
        return parsed
    if isinstance(parsed, dict):
        return {"tasks": [parsed]}
    raise PipelineStateError("source json must be an object or list")


def parse_json_task(content: str) -> dict:
    parsed = json.loads(normalize_llm_text(content))
    if isinstance(parsed, dict) and isinstance(parsed.get("tasks"), list):
        tasks = parsed["tasks"]
        if tasks:
            return tasks[0]
    if isinstance(parsed, list) and parsed:
        return parsed[0]
    if isinstance(parsed, dict):
        return parsed
    raise PipelineStateError("task json must be an object")


def merge_json_documents(parts: list[str]) -> str:
    tasks = []
    for part in parts:
        tasks.extend(parse_json_document(part).get("tasks", []))
    return json.dumps({"tasks": tasks}, ensure_ascii=False, indent=2)


def merge_source_documents(parts: list[str]) -> str:
    normalized = []
    for index, part in enumerate(parts):
        body = extract_body(part).strip()
        if not body:
            continue
        normalized.append(f"<section data-source-index=\"{index + 1}\">{body}</section>")
    if not normalized:
        return "<!doctype html><html><body><h2>Задание 1</h2><p>Источник пуст.</p></body></html>"
    return f"<!doctype html><html><body>{''.join(normalized)}</body></html>"


def extract_body(html: str) -> str:
    match = re.search(r"(?is)<body[^>]*>(.*?)</body>", html)
    if match:
        return match.group(1)
    return html


def escape_html(value: str) -> str:
    return (
        value.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


def provider_options(config: Config) -> list[dict]:
    return [
        {
            "id": "mock",
            "label": "Mock",
            "configured": True,
            "default_model": "mock",
            "models": [{"id": "mock", "label": "Mock"}],
        },
        {
            "id": "gigachat",
            "label": "GigaChat",
            "configured": bool(config.GIGACHAT_AUTH_KEY),
            "default_model": config.GIGACHAT_MODEL or "GigaChat-Pro",
            "models": [
                {"id": "GigaChat-Pro", "label": "GigaChat Pro"},
                {"id": "GigaChat-2", "label": "GigaChat 2"},
            ],
        },
        {
            "id": "openai",
            "label": "OpenAI",
            "configured": bool(config.OPENAI_API_KEY),
            "default_model": config.OPENAI_MODEL or "gpt-5.4",
            "models": [
                {"id": "gpt-5.4", "label": "GPT-5.4"},
                {"id": "gpt-5.4-mini", "label": "GPT-5.4 mini"},
                {"id": "gpt-5.4-nano", "label": "GPT-5.4 nano"},
            ],
        },
    ]


def resolve_llm_selection(
    options: dict[str, Any],
    *,
    provider_key: str,
    model_key: str,
    config: Config,
    fallback: LLMSelection | None = None,
) -> LLMSelection:
    fallback_provider = fallback.provider if fallback else config.LLM_PROVIDER
    provider = str(options.get(provider_key) or fallback_provider or "mock").strip().lower()
    if provider not in SUPPORTED_PROVIDERS:
        raise ValidationError("LLM provider is not supported.", code="invalid_llm_provider")

    fallback_model = fallback.model if fallback else default_model_for_provider(provider, config)
    raw_model = str(options.get(model_key) or "").strip()
    model = resolve_model_for_provider(raw_model, provider, config, fallback_model)
    return LLMSelection(provider=provider, model=model)


def default_model_for_provider(provider: str, config: Config) -> str:
    if provider == "gigachat":
        return config.GIGACHAT_MODEL or "GigaChat-Pro"
    if provider == "openai":
        return config.OPENAI_MODEL or "gpt-5.4"
    return "mock"


def resolve_model_for_provider(option: str, provider: str, config: Config, fallback_model: str = "") -> str:
    selected = option.strip()
    selected_lower = selected.lower()
    if selected_lower == "":
        return fallback_model or default_model_for_provider(provider, config)

    aliases = {
        "gigachat": {
            "pro": config.GIGACHAT_MODEL or "GigaChat-Pro",
            "lite": "GigaChat-2",
        },
        "openai": {
            "pro": config.OPENAI_MODEL or "gpt-5.4",
            "lite": "gpt-5.4-mini",
        },
        "mock": {
            "pro": "mock",
            "lite": "mock",
        },
    }
    if selected_lower in aliases.get(provider, {}):
        return aliases[provider][selected_lower]
    if SAFE_MODEL_RE.match(selected):
        return selected
    raise ValidationError("LLM model is invalid.", code="invalid_llm_model")


def resolve_variant_count(value: Any) -> int:
    if value in (None, ""):
        return 1
    try:
        count = int(value)
    except (TypeError, ValueError) as exc:
        raise ValidationError("Variant count must be from 1 to 10.", code="invalid_variant_count") from exc
    if count < 1 or count > 10:
        raise ValidationError("Variant count must be from 1 to 10.", code="invalid_variant_count")
    return count


def parse_results(raw: Any) -> list[dict]:
    if not raw:
        return []
    if isinstance(raw, list):
        return raw
    try:
        return json.loads(raw)
    except (TypeError, json.JSONDecodeError):
        return []


def keep_results_before(results: list[dict], step: int) -> list[dict]:
    return [result for result in results if int(result.get("step") or 0) < step]


def result_content(results: list[dict], key: str) -> str:
    for result in results:
        if result.get("key") == key:
            return str(result.get("content") or "")
    return ""


def build_pipeline_content(results: list[dict]) -> dict:
    return {
        "source_html": result_content(results, "source_html"),
        "parameters": result_content(results, "parameters"),
        "variant_html": result_content(results, "variant_html"),
        "answers_by_variant": [],
        "answers_all": "",
        "self_score": result_content(results, "self_score"),
        "steps": results,
    }


def step_result(step: int, key: str, title: str, content: str) -> dict:
    from datetime import datetime, timezone

    return {
        "step": step,
        "key": key,
        "title": title,
        "content": content,
        "created_at": datetime.now(timezone.utc).isoformat().replace("+00:00", "Z"),
    }


def normalize_llm_text(content: str) -> str:
    content = content.strip()
    if content.startswith("```"):
        lines = content.splitlines()
        if len(lines) >= 2:
            lines = lines[1:]
            if lines and lines[-1].strip() == "```":
                lines = lines[:-1]
            content = "\n".join(lines).strip()
    return content


def parse_task_parameters(content: str) -> dict:
    return {
        "task_type": extract_labeled_value(content, "Тип задания"),
        "subject": extract_labeled_value(content, "Предмет"),
        "school_class": extract_labeled_value(content, "Предполагаемый класс"),
        "difficulty": extract_labeled_value(content, "Уровень сложности задания"),
    }


def extract_labeled_value(content: str, label: str) -> str:
    for line in content.splitlines():
        line = line.strip()
        if line.lower().startswith(label.lower()):
            parts = line.split(":", 1)
            if len(parts) == 2:
                return parts[1].strip() or "*"
    return "*"


def parse_parameter_bundle(content: str) -> dict:
    try:
        bundle = json.loads(normalize_llm_text(content))
    except json.JSONDecodeError as exc:
        raise PipelineStateError(f"failed to parse step 2 parameters: {exc}") from exc
    if not bundle.get("tasks"):
        raise PipelineStateError("step 2 parameters contain no tasks")
    return bundle


def extract_task_sections(source_html: str) -> list[TaskSection]:
    try:
        tasks = parse_json_document(source_html).get("tasks", [])
    except Exception:
        tasks = []
    if tasks:
        sections = []
        for index, task in enumerate(tasks):
            task_json = json.dumps(task, ensure_ascii=False, indent=2)
            sections.append(
                TaskSection(
                    number=index + 1,
                    heading=str(task.get("title") or f"Задание {index + 1}"),
                    html=task_json,
                    start=index,
                    end=index + 1,
                )
            )
        return sections

    lower = source_html.lower()
    trailing_start = len(source_html)
    body_close = lower.rfind("</body>")
    html_close = lower.rfind("</html>")
    if body_close >= 0:
        trailing_start = body_close
    elif html_close >= 0:
        trailing_start = html_close

    starts = []
    search_from = 0
    while True:
        index = lower.find("<h2", search_from)
        if index < 0:
            break
        starts.append(index)
        search_from = index + 3
    if not starts:
        raise PipelineStateError("no task sections with <h2> were found in source html")

    sections = []
    for index, start in enumerate(starts):
        end = starts[index + 1] if index + 1 < len(starts) else trailing_start
        section_html = source_html[start:end].strip()
        sections.append(
            TaskSection(
                number=index + 1,
                heading=extract_task_heading(section_html),
                html=section_html,
                start=start,
                end=end,
            )
        )
    return sections


def extract_task_heading(section_html: str) -> str:
    lower = section_html.lower()
    start = lower.find("<h2")
    if start < 0:
        return ""
    open_end = lower.find(">", start)
    if open_end < 0:
        return ""
    close_index = lower.find("</h2>", open_end + 1)
    if close_index < 0:
        return ""
    return TAG_RE.sub(" ", section_html[open_end + 1 : close_index]).strip()


def merge_task_sections(source_html: str, sections: list[TaskSection], replacements: list[str]) -> str:
    if len(sections) != len(replacements):
        raise PipelineStateError("task replacement count does not match source task count")
    try:
        parse_json_document(source_html)
        tasks = [parse_json_task(replacement) for replacement in replacements]
        return json.dumps({"tasks": tasks}, ensure_ascii=False, indent=2)
    except json.JSONDecodeError:
        pass
    parts = [source_html[: sections[0].start]]
    for index, section in enumerate(sections):
        parts.append(replacements[index].strip())
        if index + 1 < len(sections):
            parts.append(source_html[section.end : sections[index + 1].start])
        else:
            parts.append(source_html[section.end :])
    return "".join(parts)


def contains_tag(content: str, tag: str) -> bool:
    return f"<{tag.lower()}" in content.lower()


def ensure_section_heading(original_section: str, generated: str) -> str:
    lower = original_section.lower()
    start = lower.find("<h2")
    if start < 0:
        return generated
    open_end = lower.find(">", start)
    if open_end < 0:
        return generated
    close_index = lower.find("</h2>", open_end + 1)
    if close_index < 0:
        return generated
    heading_html = original_section[start : close_index + 5]
    return heading_html + "\n" + generated.strip()


def extract_score(content: str) -> int:
    content = normalize_llm_text(content)
    try:
        parsed = json.loads(content)
        score = int(parsed.get("score"))
        if 1 <= score <= 10:
            return score
    except Exception:
        pass
    match = SCORE_RE.search(content)
    if not match:
        return 1
    score = int(match.group(1))
    return score if 1 <= score <= 10 else 1


def finished_pipeline_step(step: int) -> int:
    return EVALUATION_STEP if step == GENERATION_STEP else step


def pipeline_step_key(step: int) -> str:
    return {
        1: "source_html",
        2: "parameters",
        3: "variant_html",
        4: "self_score",
    }.get(step, "unknown")


def pipeline_step_title(step: int) -> str:
    return {
        1: "HTML исходного задания",
        2: "Параметры задания",
        3: "Новый вариант задания (HTML)",
        4: "Самооценка результата",
    }.get(step, "Шаг пайплайна")


def pipeline_task_type(step: int) -> str:
    return f"assignment_pipeline_step_{step}"


def iso(value: Any) -> str | None:
    if value is None:
        return None
    if hasattr(value, "isoformat"):
        return value.isoformat()
    return str(value)
